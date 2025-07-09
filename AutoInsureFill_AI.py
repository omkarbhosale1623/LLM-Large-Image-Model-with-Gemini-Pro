import streamlit as st
import pdfplumber
from docx import Document
import requests
import os
import re
import json
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from datetime import datetime

st.set_page_config(page_title="GLR Insurance Auto-Filler")
st.title("📄 Insurance Template Auto-Filler (via OpenRouter)")

# CONFIG
OPENROUTER_API_KEY = "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" 
LLM_MODEL = "meta-llama/llama-3.1-8b-instruct"

headers = {
    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
    "Referer": "http://localhost:8501",
    "X-Title": "GLR Insurance Auto-Filler",
    "Content-Type": "application/json"
}

def extract_template_fields(docx_file):
    """Extract placeholders like [FIELD_NAME] from the .docx template, including tables."""
    doc = Document(docx_file)
    fields = set()

    # Extract from normal paragraphs
    for para in doc.paragraphs:
        matches = re.findall(r"\[([A-Z0-9_]+)\]", para.text)
        fields.update(matches)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.findall(r"\[([A-Z0-9_]+)\]", cell.text)
                fields.update(matches)

    return sorted(list(fields))

def extract_text_from_pdf(pdf_file):
    """Extract text from a single PDF file."""
    with pdfplumber.open(pdf_file) as pdf:
        text = ""
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text

@retry(
    stop=stop_after_attempt(3),
    wait=wait_exponential(multiplier=1, min=4, max=10),
    retry=retry_if_exception_type((requests.HTTPError, requests.ConnectionError))
)
def get_extracted_data(text, fields):
    """Query LLM to extract key-value pairs for specified fields."""
    prompt = (
        f"Extract the following fields from this insurance report as key-value pairs:\n"
        f"Fields: {', '.join(fields)}\n\n"
        f"Report:\n{text}\n\n"
        f"Output ONLY valid JSON with double quotes. DO NOT include comments or explanations."
    )

    response = requests.post(
        "https://openrouter.ai/api/v1/chat/completions",
        headers=headers,
        json={
            "model": LLM_MODEL,
            "messages": [{"role": "user", "content": prompt}]
        }
    )
    response.raise_for_status()
    content = response.json()['choices'][0]['message']['content']

    # Extract JSON block using regex
    match = re.search(r'\{[\s\S]*\}', content)
    if match:
        json_str = match.group(0)

        # Remove // comments
        json_str = re.sub(r'//.*', '', json_str)

        # Remove any trailing commas before closing braces/brackets
        json_str = re.sub(r',\s*([\}\]])', r'\1', json_str)

        try:
            data = json.loads(json_str)
            data = {k: (v if v is not None else "") for k, v in data.items()}
            return data
        except json.JSONDecodeError as e:
            st.error(f"❌ Failed to decode cleaned JSON:\n{json_str}")
            raise e
    else:
        st.error(f"❌ No JSON block found in LLM response.\nFull response:\n{content}")
        raise ValueError("No JSON block found in LLM response.")

def fill_template(template_file, key_value_pairs):
    """Fill the .docx template with extracted key-value pairs."""
    doc = Document(template_file)

    for para in doc.paragraphs:
        text = para.text
        for key, value in key_value_pairs.items():
            placeholder = f"[{key}]"
            if placeholder in text:
                text = text.replace(placeholder, str(value))
        para.text = text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                for key, value in key_value_pairs.items():
                    placeholder = f"[{key}]"
                    if placeholder in text:
                        text = text.replace(placeholder, str(value))
                cell.text = text

    return doc

def main():
    with st.sidebar:
        st.header("Upload Files")
        template_file = st.file_uploader("Upload Insurance Template (.docx)", type=["docx"])
        pdf_files = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

    if st.button("Generate Filled Template"):
        if not template_file or not pdf_files:
            st.error("Please upload both a template (.docx) and at least one photo report (.pdf).")
            return

        with st.spinner("Processing files..."):
            # Extract template fields (from paragraphs and tables)
            template_fields = extract_template_fields(template_file)
            if not template_fields:
                st.error("❌ No fields found in template.")
                return
            st.write("✅ **Fields found in template:**", template_fields)

            # Process PDF reports
            all_key_value_pairs = []
            for pdf_file in pdf_files:
                st.write(f"Processing **{pdf_file.name}**...")
                report_text = extract_text_from_pdf(pdf_file)
                if report_text:
                    kv_pairs = get_extracted_data(report_text, template_fields)
                    all_key_value_pairs.append(kv_pairs)

            if not all_key_value_pairs:
                st.error("❌ No data extracted from reports.")
                return

            # Combine key-value pairs (prioritize non-empty values)
            combined_kv = {}
            for kv in all_key_value_pairs:
                for key, value in kv.items():
                    if key not in combined_kv or (combined_kv[key] == "" and value != ""):
                        combined_kv[key] = value

            # Display extracted data
            st.subheader("📄 Extracted Data (from PDFs):")
            st.json(combined_kv)

            # Fill template with combined key-value pairs
            filled_doc = fill_template(template_file, combined_kv)

            # Save the filled template
            output_dir = "task_3_output"
            os.makedirs(output_dir, exist_ok=True)
            output_path = os.path.join(output_dir, f"filled_template_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
            filled_doc.save(output_path)

            with open(output_path, "rb") as f:
                st.success("✅ Template filled successfully!")
                st.download_button(
                    label="⬇️ Download Filled Template",
                    data=f,
                    file_name="filled_template.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

if __name__ == "__main__":
    main()